from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Cm
import os
import json
# from template.config import config  # Удалить этот импорт

def replace_placeholders(doc, config):
    for paragraph in doc.paragraphs:
        # Вставка картинки подписи адвоката
        if "{lawyer_sign}" in paragraph.text:
            paragraph.clear()
            if "lawyer_sign" in config and os.path.exists(config["lawyer_sign"]):
                run = paragraph.add_run()
                run.add_picture(config["lawyer_sign"], width=Cm(3))
            else:
                paragraph.add_run(config.get("lawyer_sign", ""))
        else:
            for run in paragraph.runs:
                for key, value in config.items():
                    # Не заменяем lawyer_sign здесь, чтобы не затереть картинку
                    if key == "lawyer_sign":
                        continue
                    run.text = run.text.replace(f"{{{key}}}", str(value))
                run.font.name = "Times New Roman"
        # Добавляем разрыв страницы перед нужным абзацем
        if "4. РАЗМЕР И УСЛОВИЯ ОПЛАТЫ ТРУДА" in paragraph.text:
            paragraph_before = paragraph.insert_paragraph_before()
            paragraph_before.add_run().add_break(WD_BREAK.PAGE)
    # Заменяем также в таблицах, если они есть
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "{lawyer_sign}" in paragraph.text:
                        paragraph.clear()
                        if "lawyer_sign" in config and os.path.exists(config["lawyer_sign"]):
                            run = paragraph.add_run()
                            run.add_picture(config["lawyer_sign"], width=Cm(3))
                        else:
                            paragraph.add_run(config.get("lawyer_sign", ""))
                    else:
                        for run in paragraph.runs:
                            for key, value in config.items():
                                if key == "lawyer_sign":
                                    continue
                                run.text = run.text.replace(f"{{{key}}}", str(value))
                            run.font.name = "Times New Roman"

def generate_contract_from_template(template_path, config):
    if isinstance(config, str):
        with open(config, "r", encoding="utf-8") as f:
            config = json.load(f)
    # print(f"[DEBUG] Загрузка файла: {config}")
    doc = Document(template_path)
    replace_placeholders(doc, config)
    output_path = f"{config.get('contract_title', 'contract')}.docx"
    print(f"[DEBUG] Сохранение контракта в", config.get('contract_title', 'contract'))
    doc.save(output_path)
    return output_path